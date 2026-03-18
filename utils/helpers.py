"""
helpers.py
==========
辅助工具函数
"""

from typing import List, Dict, Optional
from io import BytesIO
from pathlib import Path

from PyPDF2 import PdfReader
from docx import Document


# ==========================================
# 文件解析函数
# ==========================================

def parse_file(uploaded_file) -> str:
    """
    解析上传文件为纯文本

    Args:
        uploaded_file: Streamlit UploadedFile 对象

    Returns:
        str: 解析后的文本内容
    """
    try:
        if uploaded_file.name.endswith('.txt'):
            return uploaded_file.read().decode("utf-8")
        if uploaded_file.name.endswith('.pdf'):
            return "".join([p.extract_text() for p in PdfReader(uploaded_file).pages])
        if uploaded_file.name.endswith('.docx'):
            return "\n".join([p.text for p in Document(uploaded_file).paragraphs])
    except:
        return ""
    return ""


def parse_file_bytes(filename: str, content: bytes) -> str:
    """
    解析文件内容（从 bytes）—— 用于从 GitHub 拉取的文件

    Args:
        filename: 文件名
        content: 文件字节内容

    Returns:
        str: 解析后的文本内容
    """
    try:
        if filename.lower().endswith('.txt'):
            return content.decode('utf-8', errors='ignore')
        elif filename.lower().endswith('.pdf'):
            if not content.startswith(b'%PDF'):
                print(f"[ERROR] 不是有效的 PDF 文件: {filename}")
                return ""
            reader = PdfReader(BytesIO(content))
            text = ""
            for page in reader.pages:
                try:
                    pt = page.extract_text()
                    if pt:
                        text += pt + "\n"
                except:
                    continue
            return text
        elif filename.lower().endswith('.docx'):
            doc = Document(BytesIO(content))
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        print(f"[ERROR] 解析 {filename} 失败: {e}")
    return ""


# ==========================================
# 报告生成函数
# ==========================================

def create_word_report(results: List[Dict]) -> BytesIO:
    """
    生成 Word 格式的批量评分报告 - 美化版

    改进：
    1. 正确处理嵌套分数结构
    2. 添加中文样式和字体
    3. 美化表格样式
    4. 添加总分和平均分

    Args:
        results: 评分结果列表

    Returns:
        BytesIO: Word 文档字节流
    """
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime

    doc = Document()

    # 设置默认字体大小
    style = doc.styles['Normal']
    style.font.size = Pt(12)

    # 标题
    title = doc.add_heading('茶评批量评分报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加生成时间和统计
    summary_para = doc.add_paragraph()
    summary_para.add_run(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').bold = True
    summary_para.add_run(f'\n评分条目数：{len(results)} 条')

    doc.add_paragraph('')
    doc.add_paragraph('=' * 40)
    doc.add_paragraph('')

    for item in results:
        # 条目标题
        heading = doc.add_heading(f'第 {item["id"]} 条茶评', level=1)

        # 原文
        para = doc.add_paragraph()
        run = para.add_run('【茶评原文】')
        run.bold = True
        run.font.size = Pt(14)
        para.add_run(f'\n{item["text"]}')

        # 提取分数（处理嵌套结构）
        scores_data = item.get('scores', {})
        if isinstance(scores_data, dict) and 'scores' in scores_data:
            scores_data = scores_data['scores']

        # 总评
        mc = item.get('master_comment', '')
        if mc:
            para = doc.add_paragraph()
            run = para.add_run('【宗师总评】')
            run.bold = True
            run.font.size = Pt(14)
            para.add_run(f'\n{mc}')

        # 计算平均分和总分
        total = 0
        count = 0
        for k, v in scores_data.items():
            if isinstance(v, dict) and 'score' in v:
                try:
                    total += int(v.get('score', 0))
                    count += 1
                except:
                    pass
        avg = total / count if count > 0 else 0

        # 添加统计信息
        stat_para = doc.add_paragraph()
        run = stat_para.add_run('【评分统计】')
        run.bold = True
        stat_para.add_run(f' 平均分：{avg:.1f}/9  |  总分：{total}/{count*9}')

        # 创建表格（4 列）
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 设置表格宽度
        for cell in table.columns[0].cells:
            cell.width = Inches(1.2)
        for cell in table.columns[1].cells:
            cell.width = Inches(0.6)
        for cell in table.columns[2].cells:
            cell.width = Inches(2.5)
        for cell in table.columns[3].cells:
            cell.width = Inches(2.5)

        # 表头
        hdr_cells = table.rows[0].cells
        headers = ['评价因子', '得分', '评语', '改进建议']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            # 设置表头样式
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)

        # 添加数据行
        for k, v in scores_data.items():
            if isinstance(v, dict) and 'score' in v:
                row_cells = table.add_row().cells
                row_cells[0].text = k
                row_cells[1].text = f"{v.get('score', '')}/9"
                row_cells[2].text = v.get('comment', '')
                row_cells[3].text = v.get('suggestion', '')

                # 设置数据行样式
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

        # 分隔线
        doc.add_paragraph('')
        doc.add_paragraph('─' * 40)
        doc.add_paragraph('')

    # 保存
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ==========================================
# 文本分割函数（批量评分专用）
# ==========================================

def split_tea_reviews(text: str, max_length: int = 2000) -> List[str]:
    """
    将长文本分割成多个独立的茶评条目 - 优化版

    改进：
    1. 增加更多分割标记支持
    2. 增大默认长度限制（500 -> 2000）
    3. 支持更多分隔符（空行、中文标点）

    Args:
        text: 原始文本
        max_length: 单条茶评的最大长度（字符数）

    Returns:
        List[str]: 分割后的茶评条目列表
    """
    # 清理文本
    text = text.strip()
    # 统一换行符
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    # 移除多余的空行
    lines = [line.strip() for line in text.split('\n') if line.strip()]

    # 如果没有内容，返回空列表
    if not lines:
        return []

    # 如果文本很短，直接返回
    if len(text) <= max_length and len(lines) <= 2:
        return [text] if text else []

    reviews = []
    current_review = ""

    # 支持的分割标记（扩展版）
    split_markers = [
        # 阿拉伯数字标记
        '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.',
        '11.', '12.', '13.', '14.', '15.', '16.', '17.', '18.', '19.', '20.',
        # 圆圈数字
        '①', '②', '③', '④', '⑤', '⑥', '⑦', '⑧', '⑨', '⑩',
        # 中文数字
        '一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、',
        '第一.', '第二.', '第三.', '第四.', '第五.', '第六.', '第七.', '第八.', '第九.', '第十.',
        # 标题标记
        '茶品', '样品', '品名', '名称', '产品', '编号',
        # 括号标记
        '【', '『', '【', '】', '』',
        # 引号标记
        '"', '"', ''', ''', '「', '」',
        # 分割线
        '---', '***', '___', '＝' * 5, '——' * 5,
    ]

    for line in lines:
        # 检查是否是新的茶评开始
        is_new_review = False
        for marker in split_markers:
            if line.startswith(marker):
                is_new_review = True
                break

        if is_new_review:
            # 保存之前的茶评
            if current_review and len(current_review) > 10:
                reviews.append(current_review.strip())
            current_review = line
        else:
            # 追加到当前茶评
            if current_review:
                current_review += " " + line
            else:
                current_review = line

        # 长度限制（增大到 2000）
        if len(current_review) >= max_length:
            reviews.append(current_review.strip())
            current_review = ""

    # 添加最后一个
    if current_review and len(current_review) > 10:
        reviews.append(current_review.strip())

    # 备用方案：如果没有成功分割（只有1条或更少），按空行分割
    if len(reviews) < 2:
        # 按段落分割（连续的空行）
        paragraphs = []
        current_para = ""

        for line in lines:
            if not line.strip():
                # 遇到空行，保存当前段落
                if current_para.strip():
                    paragraphs.append(current_para.strip())
                    current_para = ""
            else:
                current_para += " " + line if current_para else line

        # 添加最后一段
        if current_para.strip():
            paragraphs.append(current_para.strip())

        # 过滤短内容
        reviews = [p for p in paragraphs if len(p) > 10]

    # 最终过滤
    reviews = [r.strip() for r in reviews if r.strip() and len(r.strip()) > 10]

    return reviews


def parse_batch_file(uploaded_file) -> List[str]:
    """
    解析批量评分上传的文件，返回茶评条目列表

    Args:
        uploaded_file: Streamlit UploadedFile 对象

    Returns:
        List[str]: 茶评条目列表
    """
    try:
        # 解析文件内容
        text = parse_file(uploaded_file)

        if not text:
            return []

        # 分割成多个茶评
        reviews = split_tea_reviews(text)

        return reviews

    except Exception as e:
        print(f"[ERROR] 批量文件解析失败: {e}")
        return []


# ==========================================
# 模板下载函数
# ==========================================

def get_template_bytes(template_path: Path, github_sync_class) -> Optional[bytes]:
    """
    获取模板文件内容（优先本地，其次从 GitHub 下载）

    Args:
        template_path: 本地模板文件路径
        github_sync_class: GithubSync 类（用于从 GitHub 下载）

    Returns:
        Optional[bytes]: 模板文件内容，失败返回 None
    """
    if template_path.exists():
        with open(template_path, 'rb') as f:
            return f.read()
    # 尝试从 GitHub 下载
    content = github_sync_class.download_github_file("tea_backup/template.xlsx")
    if content:
        # 缓存到本地
        try:
            with open(template_path, 'wb') as f:
                f.write(content)
        except:
            pass
    return content
