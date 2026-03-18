"""
core 模块
包含后端核心逻辑
"""

import logging
import streamlit as st
from typing import List, Tuple

# 导入其他模块
from .resource_manager import ResourceManager
from .github_sync import GithubSync
from .ai_services import SelfHostedEmbedder, llm_normalize_user_input
from .scoring import run_scoring

logger = logging.getLogger(__name__)

__all__ = [
    'ResourceManager',
    'GithubSync',
    'SelfHostedEmbedder',
    'llm_normalize_user_input',
    'run_scoring',
    'bootstrap_cases',
    'load_rag_from_local',
    'load_kb_from_cache',
    'load_rag_from_github',
    'process_uploaded_files'
]


# ==========================================
# 初始化函数
# ==========================================

def bootstrap_cases(embedder):
    """
    判例库初始化（如果为空则创建默认判例）

    Args:
        embedder: 嵌入器实例
    """
    # 如果基础判例为空，可以创建一些默认判例
    if not st.session_state.basic_cases:
        pass

    # 如果进阶判例为空
    supp_idx, supp_data = st.session_state.supp_cases
    if not supp_data:
        pass


# ==========================================
# 知识库辅助函数
# ==========================================

def _get_files_metadata():
    """
    获取当前 RAG 目录中所有文件的元数据

    Returns:
        dict: {filename: {"size": int, "mtime": float}}
    """
    from config.settings import PATHS
    import os

    metadata = {}

    if not PATHS.RAG_DIR.exists():
        return metadata

    # 扫描所有支持的文件类型
    for pattern in ["*.txt", "*.pdf", "*.docx"]:
        for file_path in PATHS.RAG_DIR.glob(pattern):
            if file_path.is_file():
                stat = file_path.stat()
                metadata[file_path.name] = {
                    "size": stat.st_size,
                    "mtime": stat.st_mtime
                }

    return metadata


def _save_metadata(metadata):
    """
    保存文件元数据到 JSON 文件

    Args:
        metadata: 文件元数据字典
    """
    from config.settings import PATHS
    import json
    from datetime import datetime

    data = {
        "files": metadata,
        "last_updated": datetime.now().isoformat()
    }

    with open(PATHS.kb_metadata, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _load_cached_metadata():
    """
    从 JSON 文件加载缓存的元数据

    Returns:
        dict: 文件元数据字典，如果文件不存在返回 None
    """
    from config.settings import PATHS
    import json

    if not PATHS.kb_metadata.exists():
        return None

    try:
        with open(PATHS.kb_metadata, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return data.get("files", {})
    except Exception as e:
        logger.warning(f"加载缓存元数据失败: {e}")
        return None


def _metadata_equal(current, cached):
    """
    比较当前文件元数据和缓存的元数据是否相等

    Args:
        current: 当前文件元数据
        cached: 缓存的文件元数据

    Returns:
        bool: 如果相等返回 True
    """
    if not current or not cached:
        return False

    # 检查文件数量
    if set(current.keys()) != set(cached.keys()):
        return False

    # 检查每个文件的大小和修改时间
    for filename, current_meta in current.items():
        if filename not in cached:
            return False

        cached_meta = cached[filename]
        if (current_meta["size"] != cached_meta["size"] or
            current_meta["mtime"] != cached_meta["mtime"]):
            return False

    return True


def load_kb_from_cache():
    """
    从缓存加载知识库（快速模式）

    Returns:
        tuple: (success: bool, message: str)
    """
    from config.settings import PATHS
    import faiss
    import pickle

    try:
        # 检查缓存文件是否存在
        if not PATHS.kb_index.exists() or not PATHS.kb_chunks.exists():
            return False, "缓存文件不存在"

        logger.info("从缓存加载知识库...")

        # 加载索引
        index = faiss.read_index(str(PATHS.kb_index))

        # 加载文本块
        with open(PATHS.kb_chunks, 'rb') as f:
            chunks = pickle.load(f)

        # 更新 session_state
        st.session_state.kb = (index, chunks)

        # 加载文件列表
        kb_files = list(_get_files_metadata().keys())
        st.session_state.kb_files = kb_files

        logger.info(f"✅ 缓存加载完成：{len(chunks)} 个片段")

        return True, f"✅ 从缓存加载 {len(chunks)} 个知识片段（耗时 < 3秒）"

    except Exception as e:
        logger.error(f"缓存加载失败: {e}", exc_info=True)
        return False, f"缓存加载失败: {str(e)}"


# ==========================================
# 知识库加载函数
# ==========================================

def load_rag_from_local(embedder, force_rebuild: bool = False):
    """
    从本地 RAG 目录加载知识库（智能模式）

    功能：
    1. 检查缓存是否有效
    2. 如果缓存有效且文件未变更，直接加载缓存（快速）
    3. 如果文件有变更或强制重建，执行完整处理

    Args:
        embedder: SelfHostedEmbedder 实例
        force_rebuild: 是否强制重建（忽略缓存）

    Returns:
        tuple: (success: bool, message: str)
    """
    from config.settings import PATHS, KB_EMBEDDING_BATCH_SIZE

    try:
        # 检查本地 RAG 目录是否存在
        if not PATHS.RAG_DIR.exists():
            return False, "本地 RAG 目录不存在，请先在「知识库设计」中上传文件"

        # 获取当前文件元数据
        current_metadata = _get_files_metadata()

        if not current_metadata:
            return False, "本地 RAG 目录为空，请先在「知识库设计」中上传文件"

        # 优先尝试从缓存加载（如果不是强制重建）
        if not force_rebuild:
            if PATHS.kb_index.exists() and PATHS.kb_chunks.exists():
                cached_metadata = _load_cached_metadata()

                # 检查文件是否有变更
                if _metadata_equal(current_metadata, cached_metadata):
                    # 文件未变更，直接加载缓存
                    return load_kb_from_cache()
                else:
                    logger.info("检测到文件变更，需要重新构建知识库")
            else:
                logger.info("缓存不存在，将构建新的知识库")

        # 执行完整构建
        return _build_kb_from_scratch(embedder, current_metadata)

    except Exception as e:
        logger.error(f"知识库加载失败: {e}", exc_info=True)
        return False, f"加载失败: {str(e)}"


def _build_kb_from_scratch(embedder, metadata: dict):
    """
    从头构建知识库

    Args:
        embedder: SelfHostedEmbedder 实例
        metadata: 文件元数据

    Returns:
        tuple: (success: bool, message: str)
    """
    import os
    import pickle
    import io
    from config.settings import PATHS, KB_EMBEDDING_BATCH_SIZE
    import numpy as np

    try:
        import faiss

        logger.info("=" * 50)
        logger.info("开始构建知识库")
        logger.info("=" * 50)

        # 第一步：读取本地 RAG 文件
        logger.info("读取本地 RAG 文件...")
        local_files = []
        for filename in metadata.keys():
            file_path = PATHS.RAG_DIR / filename
            if file_path.exists():
                local_files.append(file_path)

        if not local_files:
            return False, "本地 RAG 目录为空，请先在「知识库设计」中上传文件"

        logger.info(f"找到 {len(local_files)} 个本地文件")

        # 第二步：解析文件内容
        logger.info("开始解析文件内容...")
        all_chunks = []

        for file_path in local_files:
            logger.info(f"  处理文件: {file_path.name} ({file_path.stat().st_size / 1024:.1f} KB)")

            try:
                if file_path.suffix == '.txt':
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read()
                    chunks = _chunk_text(text)
                    all_chunks.extend(chunks)

                elif file_path.suffix == '.pdf':
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text()
                    chunks = _chunk_text(text)
                    all_chunks.extend(chunks)

                elif file_path.suffix == '.docx':
                    import docx
                    doc = docx.Document(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    chunks = _chunk_text(text)
                    all_chunks.extend(chunks)

                logger.info(f"    → 提取了 {len(chunks)} 个文本片段")

            except Exception as e:
                logger.error(f"  处理文件 {file_path.name} 失败: {e}")
                continue

        if not all_chunks:
            return False, "未能从本地文件中提取到任何文本内容"

        logger.info(f"文本提取完成：总共 {len(all_chunks)} 个片段")

        # 第三步：向量化处理（使用自部署 Embedding 服务）
        batch_size = KB_EMBEDDING_BATCH_SIZE
        logger.info(f"开始向量化处理（批处理大小: {batch_size}）...")

        embeddings = []
        for i in range(0, len(all_chunks), batch_size):
            batch = all_chunks[i:i+batch_size]
            batch_embeddings = embedder.embed_texts(batch)
            embeddings.extend(batch_embeddings)
            logger.info(f"  已处理 {min(i+batch_size, len(all_chunks))}/{len(all_chunks)}")

        # 第四步：构建 FAISS 索引
        logger.info("构建 FAISS 索引...")
        embeddings_array = np.array(embeddings, dtype=np.float32)
        dimension = embeddings_array.shape[1]

        index = faiss.IndexFlatL2(dimension)
        index.add(embeddings_array)

        # 第五步：保存到本地
        logger.info("保存到本地...")
        faiss.write_index(index, str(PATHS.kb_index))

        with open(PATHS.kb_chunks, 'wb') as f:
            pickle.dump(all_chunks, f)

        # 保存元数据
        _save_metadata(metadata)

        # 第六步：更新 session_state
        st.session_state.kb = (index, all_chunks)

        kb_files = [f.name for f in local_files]
        st.session_state.kb_files = kb_files

        logger.info(f"✅ 知识库构建完成：{len(all_chunks)} 个片段")

        return True, f"✅ 成功加载 {len(local_files)} 个文件，共 {len(all_chunks)} 个知识片段"

    except Exception as e:
        logger.error(f"知识库构建失败: {e}", exc_info=True)
        return False, f"构建失败: {str(e)}"


def load_rag_from_github(embedder):
    """
    从 GitHub 加载 RAG 文件

    Args:
        embedder: SelfHostedEmbedder 实例

    Returns:
        tuple: (success: bool, message: str)
    """
    import os
    import pickle
    import io
    import numpy as np

    try:
        from core.github_sync import GithubSync
        from config.settings import PATHS
        import faiss

        # 第一步：从 GitHub 下载 RAG 文件
        logger.info("开始从 GitHub 下载 RAG 文件...")
        downloaded_files = GithubSync.pull_rag_folder(rag_folder="tea_data/RAG")

        if not downloaded_files:
            return False, "GitHub 上没有 RAG 文件，请先在「知识库设计」中上传文件"

        logger.info(f"下载完成：{len(downloaded_files)} 个文件")

        # 第二步：解析文件内容
        logger.info("开始解析文件内容...")
        all_chunks = []

        for filename, content in downloaded_files:
            logger.info(f"  处理文件: {filename} ({len(content) / 1024:.1f} KB)")

            try:
                if filename.endswith('.txt'):
                    text = content.decode('utf-8', errors='ignore')
                    chunks = _chunk_text(text)
                    all_chunks.extend(chunks)

                elif filename.endswith('.pdf'):
                    import PyPDF2
                    pdf_file = io.BytesIO(content)
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text()
                    chunks = _chunk_text(text)
                    all_chunks.extend(chunks)

                elif filename.endswith('.docx'):
                    import docx
                    doc_file = io.BytesIO(content)
                    doc = docx.Document(doc_file)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    chunks = _chunk_text(text)
                    all_chunks.extend(chunks)

                logger.info(f"    → 提取了 {len(chunks)} 个文本片段")

            except Exception as e:
                logger.error(f"  处理文件 {filename} 失败: {e}")
                continue

        if not all_chunks:
            return False, "未能从下载的文件中提取到任何文本内容"

        logger.info(f"文本提取完成：总共 {len(all_chunks)} 个片段")

        # 第三步：向量化处理
        logger.info("开始向量化处理...")

        embeddings = []
        batch_size = 25
        for i in range(0, len(all_chunks), batch_size):
            batch = all_chunks[i:i+batch_size]
            batch_embeddings = embedder.embed_texts(batch)
            embeddings.extend(batch_embeddings)
            logger.info(f"  已处理 {min(i+batch_size, len(all_chunks))}/{len(all_chunks)}")

        # 第四步：构建 FAISS 索引
        logger.info("构建 FAISS 索引...")
        embeddings_array = np.array(embeddings, dtype=np.float32)
        dimension = embeddings_array.shape[1]

        index = faiss.IndexFlatL2(dimension)
        index.add(embeddings_array)

        # 第五步：保存到本地
        logger.info("保存到本地...")
        faiss.write_index(index, str(PATHS.kb_index))

        with open(PATHS.kb_chunks, 'wb') as f:
            pickle.dump(all_chunks, f)

        # 第六步：更新 session_state
        st.session_state.kb = (index, all_chunks)

        kb_files = [fname for fname, _ in downloaded_files]
        st.session_state.kb_files = kb_files

        logger.info(f"✅ 知识库加载完成：{len(all_chunks)} 个片段")

        return True, f"✅ 成功从 GitHub 加载 {len(downloaded_files)} 个文件，共 {len(all_chunks)} 个知识片段"

    except Exception as e:
        logger.error(f"从 GitHub 加载失败: {e}", exc_info=True)
        return False, f"加载失败: {str(e)}"


def process_uploaded_files(uploaded_files, embedder, upload_to_github: bool = True) -> Tuple[bool, str, List[str]]:
    """
    处理上传的知识库文件

    Args:
        uploaded_files: Streamlit 上传的文件列表
        embedder: SelfHostedEmbedder 实例
        upload_to_github: 是否上传到 GitHub

    Returns:
        Tuple[bool, str, List[str]]: (是否成功, 消息, 处理成功的文件名列表)
    """
    import os
    import io
    import pickle
    import numpy as np

    try:
        from core.github_sync import GithubSync
        from core.resource_manager import ResourceManager
        from config.settings import PATHS
        import faiss

        # 第一步：保存文件到本地
        logger.info("开始处理上传的文件...")
        os.makedirs(PATHS.RAG_DIR, exist_ok=True)

        saved_files = []
        for uploaded_file in uploaded_files:
            file_path = PATHS.RAG_DIR / uploaded_file.name
            with open(file_path, 'wb') as f:
                f.write(uploaded_file.getbuffer())
            saved_files.append(uploaded_file.name)
            logger.info(f"  已保存: {uploaded_file.name}")

        # 第二步：解析文件内容
        logger.info("开始解析文件内容...")
        all_chunks = []

        for uploaded_file in uploaded_files:
            logger.info(f"  处理文件: {uploaded_file.name} ({uploaded_file.size / 1024:.1f} KB)")

            try:
                if uploaded_file.name.endswith('.txt'):
                    text = uploaded_file.getvalue().decode('utf-8', errors='ignore')
                    chunks = _chunk_text(text)
                    all_chunks.extend(chunks)

                elif uploaded_file.name.endswith('.pdf'):
                    import PyPDF2
                    pdf_file = io.BytesIO(uploaded_file.getvalue())
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text()
                    chunks = _chunk_text(text)
                    all_chunks.extend(chunks)

                elif uploaded_file.name.endswith('.docx'):
                    import docx
                    doc_file = io.BytesIO(uploaded_file.getvalue())
                    doc = docx.Document(doc_file)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    chunks = _chunk_text(text)
                    all_chunks.extend(chunks)

                logger.info(f"    → 提取了 {len(chunks)} 个文本片段")

            except Exception as e:
                logger.error(f"  处理文件 {uploaded_file.name} 失败: {e}")
                continue

        if not all_chunks:
            return False, "未能从上传的文件中提取到任何文本内容", []

        logger.info(f"文本提取完成：总共 {len(all_chunks)} 个片段")

        # 第三步：向量化处理
        logger.info("开始向量化处理...")

        embeddings = []
        batch_size = 25
        for i in range(0, len(all_chunks), batch_size):
            batch = all_chunks[i:i+batch_size]
            batch_embeddings = embedder.embed_texts(batch)
            embeddings.extend(batch_embeddings)
            logger.info(f"  已处理 {min(i+batch_size, len(all_chunks))}/{len(all_chunks)}")

        # 第四步：构建 FAISS 索引
        logger.info("构建 FAISS 索引...")
        embeddings_array = np.array(embeddings, dtype=np.float32)
        dimension = embeddings_array.shape[1]

        index = faiss.IndexFlatL2(dimension)
        index.add(embeddings_array)

        # 第五步：保存到本地
        logger.info("保存到本地...")
        faiss.write_index(index, str(PATHS.kb_index))

        with open(PATHS.kb_chunks, 'wb') as f:
            pickle.dump(all_chunks, f)

        # 第六步：更新 session_state
        st.session_state.kb = (index, all_chunks)
        kb_files = ResourceManager.load_kb_files()
        new_kb_files = list(set(kb_files + saved_files))
        ResourceManager.save_kb_files(new_kb_files)
        st.session_state.kb_files = new_kb_files

        logger.info(f"✅ 知识库更新完成：{len(all_chunks)} 个片段")

        # 第七步：上传到 GitHub（可选）
        github_uploaded = []
        if upload_to_github:
            logger.info("上传到 GitHub...")
            success, uploaded_names = GithubSync.add_rag_files(uploaded_files)
            if success:
                github_uploaded = uploaded_names
                logger.info(f"  已上传 {len(uploaded_names)} 个文件到 GitHub")
            else:
                logger.warning("  上传到 GitHub 失败")

        # 构建返回消息
        msg_parts = [
            f"✅ 成功处理 {len(saved_files)} 个文件",
            f"📊 提取了 {len(all_chunks)} 个知识片段",
        ]
        if github_uploaded:
            msg_parts.append(f"☁️ 已上传 {len(github_uploaded)} 个文件到 GitHub")

        return True, "\n".join(msg_parts), saved_files

    except Exception as e:
        logger.error(f"文件处理失败: {e}", exc_info=True)
        return False, f"处理失败: {str(e)}", []


def _chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """
    将文本分割成小块

    Args:
        text: 输入文本
        chunk_size: 每块大小（字符数）
        overlap: 重叠大小

    Returns:
        List[str]: 文本块列表
    """
    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = start + chunk_size
        chunk = text[start:end]

        # 只保留非空的块
        if chunk.strip():
            chunks.append(chunk.strip())

        start = end - overlap

    return chunks
